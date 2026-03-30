import sys
import re
from io import BytesIO
from pathlib import Path

import docx
import fitz  # PyMuPDF
from llama_index.core import Document

from QAWithPDF.exception import customexception
from logger import logging


def _clean_text(text: str) -> str:
    text = text.replace("\x00", " ")
    text = re.sub(r"\r\n?", "\n", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = re.sub(r"[ \t]{2,}", " ", text)
    return text.strip()


def _extract_text(file_name: str, file_bytes: bytes) -> list[Document]:
    file_type = file_name.split(".")[-1].lower()

    if file_type == "txt":
        content = _clean_text(file_bytes.decode("utf-8", errors="ignore"))
        return [Document(text=content, metadata={"filename": file_name})]

    if file_type == "pdf":
        pdf_doc = fitz.open(stream=file_bytes, filetype="pdf")
        pages: list[Document] = []
        for page_idx, page in enumerate(pdf_doc, start=1):
            raw = page.get_text() or ""
            content = _clean_text(raw)
            if content:
                pages.append(
                    Document(
                        text=content,
                        metadata={"filename": file_name, "page": page_idx, "source_kind": "page_text"},
                    )
                )
        pdf_doc.close()
        if pages:
            return pages
        return [Document(text="", metadata={"filename": file_name})]

    if file_type == "docx":
        doc = docx.Document(BytesIO(file_bytes))
        chunks: list[Document] = []

        paragraph_text = _clean_text("\n".join(para.text for para in doc.paragraphs if para.text.strip()))
        if paragraph_text:
            chunks.append(
                Document(
                    text=paragraph_text,
                    metadata={"filename": file_name, "source_kind": "paragraphs"},
                )
            )

        for table_idx, table in enumerate(doc.tables, start=1):
            rows = []
            for row in table.rows:
                row_values = [cell.text.strip() for cell in row.cells]
                if any(row_values):
                    rows.append(" | ".join(row_values))
            table_text = _clean_text("\n".join(rows))
            if table_text:
                chunks.append(
                    Document(
                        text=table_text,
                        metadata={"filename": file_name, "source_kind": "table", "table_index": table_idx},
                    )
                )

        if chunks:
            return chunks
        return [Document(text="", metadata={"filename": file_name})]

    raise ValueError("Unsupported file type. Upload .txt, .pdf, or .docx")


def load_data_from_bytes(file_name: str, file_bytes: bytes) -> list[Document]:
    try:
        logging.info("Loading file bytes for %s", file_name)
        return _extract_text(file_name=file_name, file_bytes=file_bytes)
    except Exception as e:
        logging.error("Error during document loading from bytes")
        raise customexception(e, sys)


def load_data_from_path(file_path: str | Path) -> list[Document]:
    path = Path(file_path)
    return load_data_from_bytes(file_name=path.name, file_bytes=path.read_bytes())


def load_data(uploaded_file):
    return load_data_from_bytes(file_name=uploaded_file.name, file_bytes=uploaded_file.read())
